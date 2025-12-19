import os
import json
from docx import Document
from docx.shared import Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUTPUT = os.path.join(ROOT, 'Project_Details_Readme.docx')

EXCLUDE_DIRS = {'.git', 'node_modules', 'dist', '__pycache__', '.venv'}


def build_tree(root):
    lines = []
    for dirpath, dirnames, filenames in os.walk(root):
        # compute relative path
        rel = os.path.relpath(dirpath, root)
        if any(part in EXCLUDE_DIRS for part in rel.split(os.sep)):
            # filter out excluded directories
            continue
        depth = 0 if rel == '.' else len(rel.split(os.sep))
        indent = '  ' * depth
        # directory line
        lines.append(f"{indent}- {os.path.basename(dirpath) if rel != '.' else os.path.basename(root)}")
        # files
        for f in sorted(filenames):
            if f in ('.DS_Store', 'Thumbs.db'):
                continue
            lines.append(f"{indent}  - {f}")
    return '\n'.join(lines)


def read_json_safe(path):
    try:
        with open(path, 'r', encoding='utf-8') as fh:
            return json.load(fh)
    except Exception as e:
        return None


# Build document
print('Building docx...')

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
h = doc.add_heading('Project: Judicial Chatbot', level=1)

# Overview
p = doc.add_paragraph()
p.add_run('Overview: ').bold = True
p.add_run('This repository implements a judicial assistance platform with a React frontend and an Express/MongoDB backend, integrated with OCR (Tesseract) and Gemini AI integration for assistance and summarization.')

# Folder structure
doc.add_heading('Folder Structure', level=2)

tree_text = build_tree(ROOT)
for line in tree_text.split('\n'):
    doc.add_paragraph(line, style='List Bullet')

# Tech stack
doc.add_heading('Technical Stack', level=2)

# read package.json files
client_pkg = read_json_safe(os.path.join(ROOT, 'client', 'package.json'))
server_pkg = read_json_safe(os.path.join(ROOT, 'server', 'package.json'))

p = doc.add_paragraph()
p.add_run('Frontend: ').bold = True
p.add_run('React (Vite), TailwindCSS, DaisyUI, Zustand, Axios, React Router, Recharts, React Icons\n')

p = doc.add_paragraph()
p.add_run('Backend: ').bold = True
p.add_run('Node.js, Express, MongoDB (mongoose), JWT auth, Bcrypt for password hashing, Tesseract.js for OCR, Gemini/Google Generative AI integration\n')

p = doc.add_paragraph()
p.add_run('AI/ML: ').bold = True
p.add_run('Local sentiment classifier in Sentiment-ai (scikit-learn joblib model + vectorizer), Gemini LLM integration for summarization and Q&A\n')

p = doc.add_paragraph()
p.add_run('Languages: ').bold = True
p.add_run('JavaScript (Node.js, React), Python (Sentiment-ai)\n')

# Dependencies summary
doc.add_heading('Dependencies (selected)', level=2)

if client_pkg:
    doc.add_paragraph('Client (select deps):')
    for k, v in client_pkg.get('dependencies', {}).items():
        doc.add_paragraph(f'- {k}: {v}')

if server_pkg:
    doc.add_paragraph('Server (select deps):')
    for k, v in server_pkg.get('dependencies', {}).items():
        doc.add_paragraph(f'- {k}: {v}')

# API Endpoints
doc.add_heading('API Endpoints', level=2)
endpoints = [
    ('POST', '/api/register', 'Register new user'),
    ('POST', '/api/email', 'Check email availability'),
    ('POST', '/api/password', 'Verify password and set cookie/token'),
    ('POST', '/api/user-details', 'Get user details from token'),
    ('POST', '/api/add-case', 'Add case'),
    ('POST', '/api/add-hearing', 'Add hearing'),
    ('POST', '/api/status-update', 'Update case status'),
    ('POST', '/api/ocr', 'Send image URL for OCR -> Gemini'),
    ('POST', '/api/get-case-details', 'Get case details'),
    ('POST', '/api/add-image', 'Upload/add image to case'),
]
for method, path, desc in endpoints:
    doc.add_paragraph(f'{method} {path} — {desc}')

# Data models
doc.add_heading('Data Models', level=2)
doc.add_paragraph('User Model: name, email, password (hashed), nationality, sex, cases[] (refs to Case)')
doc.add_paragraph('Case Model: author (User ref), name, description, initialResponse, hearings[], status, startdate, enddate, summary, imagedocuments[]')
doc.add_paragraph('Hearing Model: caseid (Case ref), no, date, userStatementSummary, opposingPartyStatementSummary, judgeStatementSummary, response')

# AI Components
doc.add_heading('AI Components', level=2)
doc.add_paragraph('Gemini: Uses @google/generative-ai package and GEMINI_API_KEY env var. The server uses this to answer prompts and summarize text extracted by OCR.')
doc.add_paragraph('OCR: Uses tesseract.js in the server side; recognized text is sent to Gemini for analysis.')
doc.add_paragraph('Sentiment-ai: A separate Python subfolder (Sentiment-ai) with mewo.py that demonstrates using a trained model (trained_model.sav) and vectorizer (vectorizer.pkl) to classify sentiment. Requires joblib and nltk to run.')

# Environment variables
doc.add_heading('Environment Variables', level=2)
doc.add_paragraph('Common .env variables used:')
doc.add_paragraph('- MONGODB_URI: MongoDB connection string')
doc.add_paragraph('- JWT_SECRET_KEY: Secret for signing JWT tokens')
doc.add_paragraph('- FRONTEND_URL: Frontend origin for CORS')
doc.add_paragraph('- GEMINI_API_KEY: API key for Google Generative AI (Gemini)')

# Setup & Run
doc.add_heading('Setup & Run', level=2)
doc.add_paragraph('1. Install server dependencies: from /server run `npm install`')
doc.add_paragraph('2. Install client dependencies: from /client run `npm install`')
doc.add_paragraph('3. Create a .env in /server with required variables (MONGODB_URI, JWT_SECRET_KEY, FRONTEND_URL, GEMINI_API_KEY)')
doc.add_paragraph('4. Run server in dev: `npm run dev` (server)')
doc.add_paragraph('5. Run client in dev: `npm run dev --prefix client`')
doc.add_paragraph('6. Build client: `npm run build --prefix client` and server serves static files from client/dist for production')

# Notes & Gotchas
doc.add_heading('Notes & Gotchas', level=2)
doc.add_paragraph('- The server sets process.env.NODE_TLS_REJECT_UNAUTHORIZED = "0" which accepts self-signed certs; be cautious in production.')
doc.add_paragraph('- The OCR flow expects an image URL; ensure images are accessible by the server.')
doc.add_paragraph('- Sentiment-ai uses joblib-pickled model files; these are included under Sentiment-ai folder.')

# Contribution and License
doc.add_heading('Contributing & License', level=2)
doc.add_paragraph('Add contribution guidelines and license information here.')

# Save
print('Saving to:', OUTPUT)
doc.save(OUTPUT)
print('Done')
